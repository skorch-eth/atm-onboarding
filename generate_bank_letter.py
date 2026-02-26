"""Banking Relationship Letter generator — pure Python (replaces generate_bank_letter.js)"""
import sys, json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_run(para, text, bold=False, size=11):
    run = para.add_run(text)
    run.bold = bold
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    return run


def add_para(doc, text='', bold=False, size=11, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        add_run(p, text, bold=bold, size=size)
    return p


def add_line(doc, label, value, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p, label, bold=True, size=size)
    add_run(p, f'  {value}', size=size)


def add_blank_line(doc, label, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p, label, bold=True, size=size)
    add_run(p, '  _________________________________', size=size)


def set_bottom_border(para):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_all_borders(para, color='CCCCCC'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '1')
        el.set(qn('w:color'), color)
        pBdr.append(el)
    pPr.append(pBdr)


def generate(data, output_path):
    m = data['merchant']
    b = data.get('bank', {})
    today = data.get('date', '')
    company_name = m['company_name']
    bank_name = b.get('bank_name') or '[Bank Name]'
    address = f"{m['company_address']}, {m['company_city']}, {m['company_state']} {m['company_zip']}"
    fein = m.get('fein') or '[EIN]'

    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    add_para(doc, today, size=11, space_after=10)
    add_para(doc, 'To Whom It May Concern', bold=True, size=11, space_after=2)
    add_para(doc, bank_name, size=11, space_after=10)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    add_run(p, 'RE: ', bold=True, size=11)
    add_run(p, f'Banking Relationship Letter — {company_name}', size=11)

    # Divider
    div = add_para(doc, space_after=8)
    set_bottom_border(div)

    add_para(doc, 'Bank Confirmation', bold=True, size=11, space_after=6)
    add_para(doc,
        f'The undersigned bank officer hereby confirms that {company_name} has established a '
        f'business banking relationship with {bank_name} and that the account information below is accurate.',
        size=11, space_after=10)

    add_line(doc, 'Bank Name:', bank_name)
    add_line(doc, 'Account Holder:', company_name)
    add_line(doc, 'Principal Address:', address)
    add_line(doc, 'Federal Employer Identification Number (FEIN):', fein)
    add_line(doc, 'Account Type:', 'Business Checking')
    add_blank_line(doc, 'Routing Number:')
    add_blank_line(doc, 'Account Number:')
    add_line(doc, 'Account Status:', 'Active  /  In Good Standing')

    add_para(doc, space_after=10)
    add_para(doc, 'Bank Officer Signature:', bold=True, size=11, space_after=4)

    # Signature line
    sig = add_para(doc, ' ', size=22, space_after=6)
    set_bottom_border(sig)

    # Printed name / title row
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p, 'Printed Name:  ___________________________', size=11)
    add_run(p, '        ', size=11)
    add_run(p, 'Title:  ___________________________', size=11)

    # Date / branch row
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    add_run(p, 'Date:  ___________________________', size=11)
    add_run(p, '        ', size=11)
    add_run(p, 'Branch:  ___________________________', size=11)

    add_para(doc, 'Bank Stamp / Seal:', bold=True, size=11, space_after=4)

    # Stamp box
    stamp = add_para(doc, ' ', size=36, space_after=0)
    set_all_borders(stamp, 'CCCCCC')

    doc.save(output_path)
    print(f'OK: {output_path}')


if __name__ == '__main__':
    data = json.loads(sys.argv[1])
    output_path = sys.argv[2]
    generate(data, output_path)
