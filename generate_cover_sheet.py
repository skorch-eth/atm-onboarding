"""Cover Sheet generator — pure Python (replaces generate_cover_sheet.js)"""
import sys, json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BLUE     = '1A3C5E'
MID_BLUE = '2C5F8A'
LT_BLUE  = 'EAF1F8'
GRAY     = '666666'
RULE     = 'BBCFE0'
GREEN    = '1A6B3A'
LT_GREEN = 'EAF7EF'


def rgb(h):
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def cell_borders(cell, color=RULE):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_run(para, text, bold=False, size=8, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    return run


def section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text, bold=True, size=8, color=MID_BLUE)


def info_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    for label, value in rows:
        row = table.add_row()
        c0, c1 = row.cells[0], row.cells[1]
        set_cell_bg(c0, LT_BLUE)
        cell_borders(c0)
        cell_borders(c1)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        add_run(p0, label, bold=True, size=8, color=BLUE)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        add_run(p1, value or '—', size=8, color='000000' if value else GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def step_table(doc, steps):
    """steps: list of (num, title, items)"""
    table = doc.add_table(rows=0, cols=3)
    table.style = 'Table Grid'
    for num, title, items in steps:
        row = table.add_row()
        c0, c1, c2 = row.cells[0], row.cells[1], row.cells[2]

        # Number cell
        set_cell_bg(c0, BLUE)
        cell_borders(c0, BLUE)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_after = Pt(0)
        add_run(p0, str(num), bold=True, size=13, color='FFFFFF')

        # Content cell
        cell_borders(c1)
        p_title = c1.paragraphs[0]
        p_title.paragraph_format.space_after = Pt(2)
        add_run(p_title, title, bold=True, size=8, color=BLUE)
        for item in items:
            p_item = c1.add_paragraph()
            p_item.paragraph_format.space_after = Pt(1)
            add_run(p_item, f'  {item}', size=8, color='333333')

        # Status cell
        set_cell_bg(c2, 'F7FAFD')
        cell_borders(c2)
        p_status = c2.paragraphs[0]
        p_status.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_status.paragraph_format.space_after = Pt(4)
        add_run(p_status, '□ Pending', bold=True, size=8, color=GRAY)
        p_date = c2.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_date.paragraph_format.space_after = Pt(0)
        add_run(p_date, 'Date: __________', size=8, color=GRAY)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def checklist_table(doc, docs):
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    for name in docs:
        row = table.add_row()
        c0, c1 = row.cells[0], row.cells[1]
        set_cell_bg(c0, LT_GREEN)
        cell_borders(c0)
        cell_borders(c1)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_after = Pt(0)
        add_run(p0, '☑', bold=True, size=8, color=GREEN)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        add_run(p1, name, size=8)


def generate(data, output_path):
    m = data['merchant']
    b = data.get('bank', {})
    today = data.get('date', '')
    company_name = m['company_name']
    owner_name = m.get('entity_creator_name', '')
    state = m.get('company_state', '')
    address = f"{m.get('company_address','')}, {m.get('company_city','')}, {state} {m.get('company_zip','')}"
    phone = m.get('location_phone', '')
    email = m.get('email', '')


    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.625)
        section.right_margin = Inches(0.625)

    # Header bar
    hdr = doc.add_paragraph()
    hdr.paragraph_format.space_before = Pt(0)
    hdr.paragraph_format.space_after = Pt(2)
    pPr = hdr._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), BLUE)
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    add_run(hdr, '  ATM MERCHANT ONBOARDING  |  PACKET COVER SHEET', bold=True, size=12, color='FFFFFF')

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    section_heading(doc, 'MERCHANT INFORMATION')
    info_table(doc, [
        ('Legal Entity Name', company_name),
        ('Owner / Managing Member', owner_name),
        ('State of Formation', state),
        ('Principal Address', address),
        ('Phone', phone),
        ('Email', email),
        ('Date Prepared', today),
    ])

    section_heading(doc, 'ONBOARDING STATUS')
    step_table(doc, [
        (1, 'File for LLC', [
            '□ Articles of Organization submitted',
            '□ State filing fee paid',
            '□ Approval documents received',
            '□ Operating Agreement drafted',
        ]),
        (2, 'Obtain EIN', [
            '□ IRS online application completed (irs.gov/ein)',
            '□ CP 575 confirmation letter downloaded',
            '□ EIN recorded: ____________________________',
        ]),
        (3, 'Open Bank Account', [
            '□ Business checking account opened',
            '□ Routing number obtained',
            '□ Account number obtained',
            '□ Banking Relationship Letter signed by banker',
        ]),
        (4, 'Emailing the Bank Letter and EIN Letter', [
            '□ Email the Bank Letter',
            '□ Email the EIN Letter',
        ]),
    ])

    doc.save(output_path)
    print(f'OK: {output_path}')


if __name__ == '__main__':
    data = json.loads(sys.argv[1])
    output_path = sys.argv[2]
    generate(data, output_path)
