"""EIN Application Answer Sheet generator — pure Python (replaces generate_ein_sheet.js)"""
import sys, json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.shared import OxmlElement as OE


BLUE      = '2C5282'
LT_BLUE   = 'EBF5FB'
RED       = 'C0392B'
LT_RED    = 'FDEDEC'
ORANGE    = 'E67E22'
LT_ORANGE = 'FEF9E7'
GRAY      = '666666'


def rgb(hex_str):
    r, g, b = int(hex_str[0:2],16), int(hex_str[2:4],16), int(hex_str[4:6],16)
    return RGBColor(r, g, b)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def cell_borders(cell, color='CCCCCC'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_run(para, text, bold=False, italic=False, size=11, color=None, font='Arial'):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    return run


def heading_para(doc, text, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = add_run(p, text, bold=True, size=13, color=color)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def step_table(doc, rows_data):
    """rows_data: list of (step_num, question, answer, note)"""
    table = doc.add_table(rows=0, cols=3)
    table.style = 'Table Grid'
    # Header row
    hdr = table.add_row()
    for i, (txt, w) in enumerate([('#', 640), ('IRS Question', 3560), ('Your Answer (copy exactly)', 5160)]):
        cell = hdr.cells[i]
        cell.width = Pt(w)
        set_cell_bg(cell, '2C5282')
        cell_borders(cell, '2C5282')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, txt, bold=True, size=10, color='FFFFFF')

    for (step_num, question, answer, note) in rows_data:
        row = table.add_row()
        # Step num
        c0 = row.cells[0]
        set_cell_bg(c0, 'EBF5FB')
        cell_borders(c0)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p0, str(step_num), bold=True, size=11, color=BLUE)

        # Question
        c1 = row.cells[1]
        cell_borders(c1)
        add_run(c1.paragraphs[0], question, size=11)

        # Answer + note
        c2 = row.cells[2]
        set_cell_bg(c2, 'FDFEFE')
        cell_borders(c2)
        p2 = c2.paragraphs[0]
        add_run(p2, answer, bold=True, size=11, color='1A5276')
        if note:
            p_note = c2.add_paragraph()
            add_run(p_note, note, italic=True, size=9, color='666666')

    doc.add_paragraph()  # spacer


def generate(data, output_path):
    m = data['merchant']
    today = data.get('date', '')
    company_name = m['company_name']
    signer_name = m.get('entity_creator_name', '')
    signer_title = m.get('title', 'Managing Member')
    address1 = m.get('company_address', '')
    city = m.get('company_city', '')
    state = m.get('company_state', '')
    zip_ = m.get('company_zip', '')
    phone = m.get('location_phone', '')
    dba = m.get('dba_name', '')

    phone_digits = ''.join(c for c in phone if c.isdigit())
    if len(phone_digits) == 10:
        phone_fmt = f"{phone_digits[:3]}-{phone_digits[3:6]}-{phone_digits[6:]}"
    else:
        phone_fmt = phone

    name_parts = signer_name.split(' ', 1)
    first_name = name_parts[0]
    last_name = name_parts[1] if len(name_parts) > 1 else ''

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Title
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p, 'EIN Application Answer Sheet', bold=True, size=18, color=BLUE)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    add_run(p2, company_name, bold=True, size=13)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(10)
    add_run(p3, f'Prepared: {today}', size=10, color=GRAY)

    # Instructions box (simple bordered paragraphs)
    instructions = [
        ('⚠  HOW TO USE THIS SHEET', True, ORANGE),
        ('1. Open: https://irs.gov/ein  (search "IRS EIN online application")', False, '000000'),
        ('2. Click "Apply Online Now" — do NOT close the browser during the session (30-min timeout)', False, '000000'),
        ('3. Follow the steps below — answer each screen exactly as shown in blue', False, '000000'),
        ('4. At the end, download and save the EIN Confirmation Letter (CP 575) immediately', False, '000000'),
        ('5. Note the EIN in the box at the bottom of this sheet', False, '000000'),
    ]
    for text, bold, color in instructions:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_run(p, f'  {text}', bold=bold, size=11, color=color)
    doc.add_paragraph()

    # Section 1
    heading_para(doc, 'SECTION 1 — Entity Type')
    p = doc.add_paragraph()
    add_run(p, 'The first screens ask what kind of entity is applying. Select exactly as shown.', italic=True, size=11, color='444444')
    step_table(doc, [
        (1, 'What type of legal structure is your business?', 'Limited Liability Company (LLC)', ''),
        (2, 'How many members does the LLC have?', '1', 'Even if multiple — select 1 if a single member owns the LLC for tax purposes'),
        (3, 'Which state is the LLC organized in?', state, ''),
    ])

    # Section 2
    heading_para(doc, 'SECTION 2 — Responsible Party')
    p = doc.add_paragraph()
    add_run(p, "The \"responsible party\" is the natural person who owns or controls the LLC. This person's SSN is required.", italic=True, size=11, color='444444')
    step_table(doc, [
        (4, 'First name', first_name, ''),
        (5, 'Last name', last_name, ''),
        (6, 'SSN or ITIN of responsible party', '***-**-****', '⚠ Fill in your actual SSN here — do not share this sheet after completing'),
        (7, 'Title / role', signer_title, ''),
    ])

    # Section 3
    heading_para(doc, 'SECTION 3 — Business Information')
    p = doc.add_paragraph()
    add_run(p, 'These screens ask about the business itself.', italic=True, size=11, color='444444')
    dba_answer = dba if dba and dba != company_name else 'Leave blank'
    dba_note = 'Enter DBA name if applicable' if dba and dba != company_name else 'Only fill if different from legal name'
    step_table(doc, [
        (8,  'Legal name of business', company_name, ''),
        (9,  'Trade name / DBA (if different)', dba_answer, dba_note),
        (10, 'Business mailing address — Street', address1, ''),
        (11, 'City', city, ''),
        (12, 'State', state, ''),
        (13, 'ZIP code', zip_, ''),
        (14, 'County', '', 'Enter the county where the business is located (e.g., Miami-Dade)'),
        (15, 'Business phone number', phone_fmt or '(enter business phone)', ''),
    ])

    # Section 4
    heading_para(doc, "SECTION 4 — Why You're Applying")
    p = doc.add_paragraph()
    add_run(p, 'The IRS asks the reason for applying for an EIN.', italic=True, size=11, color='444444')
    step_table(doc, [
        (16, 'Why are you applying for an EIN?', 'Started a new business', ''),
        (17, 'Date business started or was acquired', today, "Use today's date or the LLC formation date if already filed"),
        (18, 'Closing month of your accounting year', 'December', ''),
        (19, 'Highest number of employees expected in 12 months', '0', "If you don't plan to have W-2 employees, enter 0"),
        (20, 'Does the business have, or expect to have, employees?', 'No', "Select No if no W-2 payroll — you can always get a payroll EIN later"),
        (21, 'Principal activity of your business', 'Finance and Insurance', 'ATM operation falls under financial services'),
        (22, 'Specific product or service', 'ATM operation and management', ''),
    ])

    # EIN record box
    heading_para(doc, 'RECORD YOUR EIN HERE', RED)
    p = doc.add_paragraph()
    add_run(p, 'After submitting, the IRS shows your EIN immediately. Write it here before closing the window.', italic=True, size=11, color=GRAY)

    p_ein = doc.add_paragraph()
    p_ein.paragraph_format.space_after = Pt(6)
    p_ein.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_ein, f'EIN for {company_name}:   ', bold=True, size=12)
    add_run(p_ein, '__ __ - __ __ __ __ __ __ __', bold=True, size=14, color=RED, font='Courier New')

    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_note, 'Download the CP 575 confirmation letter and send a copy to your ISO', italic=True, size=10, color=GRAY)

    doc.save(output_path)
    print(f'OK: {output_path}')


if __name__ == '__main__':
    data = json.loads(sys.argv[1])
    output_path = sys.argv[2]
    generate(data, output_path)
